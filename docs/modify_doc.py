from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('功能测试文档.docx')


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    return table


def make_heading_element(text, level):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading%d' % level)
    pPr.append(pStyle)
    p.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    p.append(run)
    return p


def make_paragraph_element(text):
    p = OxmlElement('w:p')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    p.append(run)
    return p


def make_empty_paragraph():
    return OxmlElement('w:p')


# ========== Insert "测试设计方法选择" before "一、用户端功能测试" ==========
target_p = doc.paragraphs[3]._element

elements = []
elements.append(make_heading_element('〇、测试设计方法选择', 1))
elements.append(make_paragraph_element(
    '本系统功能测试采用以下测试设计方法，根据不同功能模块的特点选择合适的方法或方法组合：'))
elements.append(make_empty_paragraph())

for elem in reversed(elements):
    target_p.addprevious(elem)

# ========== Add method tables and supplementary cases at end ==========
UH = ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过']

doc.add_paragraph('')
add_table(doc, ['序号', '设计方法', '适用场景', '本系统应用模块'], [
    ['1', '等价类划分法', '输入条件有明确取值范围或集合', '用户注册、登录、评价评分、数量输入'],
    ['2', '边界值分析法', '输入/输出存在边界条件', '价格区间、库存数量、密码长度、地址数量'],
    ['3', '判定表法', '多个条件组合影响结果', '商家审核状态、订单状态流转、用户权限'],
    ['4', '状态迁移法', '系统存在明确的状态转换', '订单生命周期、商家审核流程、图书上下架'],
    ['5', '场景法', '涉及多个步骤的业务流程', '购物流程、下单支付流程、评价流程'],
    ['6', '错误推测法', '基于经验推测可能的错误', '接口异常、并发操作、数据越权'],
])

doc.add_paragraph('')
doc.add_heading('1. 等价类划分法', level=2)
doc.add_paragraph('原理：将输入数据划分为若干等价类，从每个等价类中选取一个代表性数据作为测试用例。')
doc.add_paragraph('应用示例 — 用户注册：')
add_table(doc, ['等价类类型', '输入项', '有效等价类', '无效等价类'], [
    ['用户名', 'username', '4-20位字母数字组合', '空、<4位、>20位、特殊字符'],
    ['密码', 'password', '6-20位，含字母和数字', '空、<6位、>20位、纯字母、纯数字'],
    ['确认密码', 'confirmPassword', '与密码一致', '与密码不一致'],
])

doc.add_paragraph('')
doc.add_heading('2. 边界值分析法', level=2)
doc.add_paragraph('原理：选取输入/输出范围的边界值及其邻近值作为测试用例。')
doc.add_paragraph('应用示例 — 密码长度（6-20位）：')
add_table(doc, ['边界点', '测试值', '预期结果'], [
    ['最小长度-1', '5位密码', '注册失败，提示密码长度不足'],
    ['最小长度', '6位密码', '注册成功'],
    ['最大长度', '20位密码', '注册成功'],
    ['最大长度+1', '21位密码', '注册失败，提示密码过长'],
])

doc.add_paragraph('')
doc.add_heading('3. 判定表法', level=2)
doc.add_paragraph('原理：列出所有条件和动作，分析条件组合与动作之间的对应关系。')
doc.add_paragraph('应用示例 — 商家登录判定表：')
add_table(doc, ['条件\\规则', 'R1', 'R2', 'R3', 'R4', 'R5'], [
    ['账号存在', 'Y', 'Y', 'Y', 'Y', 'N'],
    ['密码正确', 'Y', 'Y', 'Y', 'N', '-'],
    ['审核状态=待审核', 'Y', 'N', 'N', '-', '-'],
    ['审核状态=正常', 'N', 'Y', 'N', '-', '-'],
    ['审核状态=已拒绝', 'N', 'N', 'Y', '-', '-'],
    ['-> 登录成功', '', 'V', '', '', ''],
    ['-> 提示审核中', 'V', '', '', '', ''],
    ['-> 提示已拒绝', '', '', 'V', '', ''],
    ['-> 提示密码错误', '', '', '', 'V', ''],
    ['-> 提示账号不存在', '', '', '', '', 'V'],
])

doc.add_paragraph('')
doc.add_heading('4. 状态迁移法', level=2)
doc.add_paragraph('原理：分析系统对象的状态及其迁移条件，覆盖所有合法和非法状态迁移路径。')
doc.add_paragraph('订单状态流转：待付款 -> 待发货 -> 已发货 -> 已完成（待付款/待发货可 -> 已取消）')
doc.add_paragraph('状态迁移测试用例：')
add_table(doc, ['编号', '迁移路径', '测试步骤', '预期结果'], [
    ['ST-01', '初始->待付款', '提交订单', '订单创建成功，状态为待付款'],
    ['ST-02', '待付款->待发货', '模拟支付', '支付成功，状态变为待发货'],
    ['ST-03', '待付款->已取消', '取消订单', '订单取消，库存恢复'],
    ['ST-04', '待发货->已发货', '商家发货', '状态变为已发货'],
    ['ST-05', '已发货->已完成', '确认收货', '状态变为已完成'],
    ['ST-06', '已发货->待付款（非法）', '尝试回退状态', '操作失败，不允许回退'],
    ['ST-07', '已完成->待发货（非法）', '尝试回退状态', '操作失败，不允许回退'],
    ['ST-08', '已取消->待付款（非法）', '尝试重新激活', '操作失败，不允许激活'],
])

doc.add_paragraph('')
doc.add_heading('5. 场景法', level=2)
doc.add_paragraph('原理：模拟用户实际操作场景，覆盖基本流、备选流和异常流。')
doc.add_paragraph('应用示例 — 完整购物流景：')
doc.add_paragraph('基本流：用户登录 -> 浏览图书 -> 加入购物车 -> 去结算 -> 选择地址 -> 提交订单 -> 模拟支付 -> 商家发货 -> 确认收货 -> 提交评价 -> 商家回复')
doc.add_paragraph('备选流1：从详情页直接购买（跳过购物车）')
doc.add_paragraph('备选流2：修改购物车商品数量后结算')
doc.add_paragraph('备选流3：只选择部分商品结算')
doc.add_paragraph('异常流1：库存不足时下单')
doc.add_paragraph('异常流2：商品已下架时下单')

doc.add_paragraph('')
doc.add_heading('6. 错误推测法', level=2)
doc.add_paragraph('原理：基于经验和直觉，推测程序中可能存在的错误，设计针对性测试用例。')
add_table(doc, ['序号', '推测错误', '测试用例'], [
    ['1', '并发下单导致超卖', '同一用户快速多次点击提交订单'],
    ['2', '价格篡改', '前端修改价格参数后提交订单'],
    ['3', '越权访问', '用户A查看/操作用户B的订单'],
    ['4', 'XSS注入', '在评价内容中输入<script>标签'],
    ['5', 'SQL注入', "在搜索框输入 ' OR 1=1 --"],
    ['6', '重复提交', '快速连续点击提交按钮'],
])

# ========== 补充测试用例 ==========
doc.add_paragraph('')
doc.add_heading('六、补充测试用例', level=1)

doc.add_heading('6.1 注册表单验证补充（等价类 + 边界值）', level=2)
add_table(doc, UH, [
    ['U-055', '注册-用户名为空', '不输入用户名直接提交', '提示请输入用户名', '', ''],
    ['U-056', '注册-用户名3位', '输入3位用户名', '提示用户名长度不足', '', ''],
    ['U-057', '注册-用户名4位', '输入4位用户名', '注册成功', '', ''],
    ['U-058', '注册-密码为空', '不输入密码直接提交', '提示请输入密码', '', ''],
    ['U-059', '注册-密码5位', '输入5位密码', '提示密码长度不足', '', ''],
    ['U-060', '注册-密码与确认密码不一致', '输入不同密码', '提示两次密码不一致', '', ''],
    ['U-061', '注册-用户名含特殊字符', '输入user@#name', '提示用户名格式不正确', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.2 搜索功能补充（等价类 + 错误推测）', level=2)
add_table(doc, UH, [
    ['U-062', '搜索-空关键词', '不输入内容直接搜索', '显示全部图书或提示输入关键词', '', ''],
    ['U-063', '搜索-超长关键词', '输入200+字符搜索', '正常处理，不报错', '', ''],
    ['U-064', '搜索-XSS攻击', '输入<script>alert(1)</script>', '正常显示，不执行脚本', '', ''],
    ['U-065', '搜索-SQL注入', "输入 ' OR 1=1 --", '正常返回空结果', '', ''],
    ['U-066', '搜索-无结果', '搜索不存在的书名', '显示"暂无搜索结果"', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.3 购物车补充（边界值 + 场景法）', level=2)
add_table(doc, UH, [
    ['U-067', '购物车-数量为0', '将数量修改为0', '提示数量不能为0或自动删除', '', ''],
    ['U-068', '购物车-数量为负数', '将数量修改为-1', '提示数量不合法', '', ''],
    ['U-069', '购物车-数量超库存', '将数量修改为库存+1', '提示超出库存限制', '', ''],
    ['U-070', '购物车-全选', '点击全选按钮', '所有商品被选中，合计金额更新', '', ''],
    ['U-071', '购物车-取消全选', '取消全选按钮', '所有商品取消选中，合计为0', '', ''],
    ['U-072', '购物车-重复添加', '对同一图书多次加入购物车', '数量累加而非新增一条', '', ''],
    ['U-073', '购物车-商品下架', '加入购物车后商家下架该书', '结算时提示商品已下架', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.4 订单流程补充（状态迁移 + 判定表）', level=2)
add_table(doc, UH, [
    ['U-074', '下单-库存扣减', '购买3本库存为5的书', '下单后库存变为2', '', ''],
    ['U-075', '取消订单-库存恢复', '取消待付款订单', '库存恢复原值', '', ''],
    ['U-076', '重复支付', '对已支付订单再次点击支付', '提示订单已支付', '', ''],
    ['U-077', '已取消订单支付', '对已取消订单点击支付', '提示订单已取消，无法支付', '', ''],
    ['U-078', '已完成订单取消', '对已完成订单点击取消', '不显示取消按钮或提示不可取消', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.5 收货地址补充（边界值 + 判定表）', level=2)
add_table(doc, UH, [
    ['U-079', '地址-上限20个', '添加第21个地址', '提示地址数量已达上限', '', ''],
    ['U-080', '地址-第20个', '添加第20个地址', '添加成功', '', ''],
    ['U-081', '地址-收件人为空', '不填收件人直接保存', '提示请输入收件人', '', ''],
    ['U-082', '地址-电话格式错误', '输入非11位手机号', '提示手机号格式不正确', '', ''],
    ['U-083', '地址-删除默认地址', '删除已设为默认的地址', '自动将另一个地址设为默认', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.6 评价功能补充（等价类 + 错误推测）', level=2)
add_table(doc, UH, [
    ['U-084', '评价-评分为0星', '选择0星提交', '提示请选择评分', '', ''],
    ['U-085', '评价-内容为空', '不输入评价内容直接提交', '提示请输入评价内容或允许提交', '', ''],
    ['U-086', '评价-内容超长', '输入5000+字评价', '提示内容过长或正常截断', '', ''],
    ['U-087', '评价-XSS攻击', '输入<img onerror=alert(1) src=x>', '正常显示文本，不执行脚本', '', ''],
    ['U-088', '评价-未登录评价', '未登录状态访问评价页', '跳转登录页', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.7 个人中心补充（等价类 + 边界值）', level=2)
add_table(doc, UH, [
    ['U-089', '修改密码-新旧相同', '输入与原密码相同的新密码', '提示新密码不能与原密码相同', '', ''],
    ['U-090', '个人信息-邮箱格式错误', '输入abc@格式邮箱', '提示邮箱格式不正确', '', ''],
    ['U-091', '个人信息-手机号格式错误', '输入10位手机号', '提示手机号格式不正确', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.8 商家图书管理补充（等价类 + 边界值）', level=2)
add_table(doc, UH, [
    ['S-027', '新增图书-价格为0', '输入价格0', '提示价格不合法', '', ''],
    ['S-028', '新增图书-价格为负数', '输入价格-10', '提示价格不合法', '', ''],
    ['S-029', '新增图书-库存为负数', '输入库存-1', '提示库存不合法', '', ''],
    ['S-030', '新增图书-书名为空', '不输入书名', '提示请输入书名', '', ''],
    ['S-031', '新增图书-ISBN重复', '输入已存在的ISBN', '提示ISBN已存在', '', ''],
    ['S-032', '新增图书-封面过大', '上传10MB+图片', '提示图片大小超限', '', ''],
    ['S-033', '新增图书-非图片格式', '上传.txt文件作为封面', '提示格式不支持', '', ''],
    ['S-034', '下架图书-有关联订单', '下架已有待发货订单的图书', '下架成功，不影响已有订单', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.9 商家订单与评价补充（状态迁移 + 错误推测）', level=2)
add_table(doc, UH, [
    ['S-035', '发货-待付款订单', '对待付款订单点击发货', '不显示发货按钮或提示不可发货', '', ''],
    ['S-036', '发货-已发货订单', '对已发货订单再次发货', '不显示发货按钮或提示已发货', '', ''],
    ['S-037', '查看其他商家订单', '尝试查看不属于自己的订单', '返回404或无数据', '', ''],
    ['S-038', '回复-内容为空', '不输入内容直接提交', '提示请输入回复内容', '', ''],
    ['S-039', '回复-其他商家评价', '尝试回复其他店铺的评价', '返回403或无权限提示', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.10 管理员端补充（判定表 + 状态迁移）', level=2)
add_table(doc, UH, [
    ['A-030', '审核-已通过商家', '对已通过商家再次审核', '不显示审核按钮或提示已审核', '', ''],
    ['A-031', '审核-拒绝原因为空', '拒绝时不填原因', '提示请输入拒绝原因', '', ''],
    ['A-032', '禁用-待审核商家', '对待审核商家点击禁用', '不可操作或提示需先审核', '', ''],
    ['A-033', '启用-已拒绝商家', '对已拒绝商家点击启用', '不可操作或提示需重新审核', '', ''],
    ['A-034', '禁用-已禁用用户', '对已禁用用户再次禁用', '提示已禁用或按钮不可用', '', ''],
    ['A-035', '新增分类-名称为空', '不输入名称直接保存', '提示请输入分类名称', '', ''],
    ['A-036', '新增分类-名称重复', '输入已存在的分类名', '提示分类名已存在', '', ''],
    ['A-037', '删除分类-有上架图书', '删除有上架图书的分类', '提示该分类下有图书，无法删除', '', ''],
])

doc.add_paragraph('')
doc.add_heading('6.11 接口安全补充（错误推测 + 边界值）', level=2)
add_table(doc, UH, [
    ['E-008', 'Token伪造', '使用伪造的JWT Token访问接口', '返回401', '', ''],
    ['E-009', 'Token过期', '使用过期Token访问接口', '返回401，提示重新登录', '', ''],
    ['E-010', '参数篡改-价格', '修改请求中的价格参数', '后端忽略前端价格，按数据库价格计算', '', ''],
    ['E-011', '参数篡改-用户ID', '修改请求中的userId', '返回403或使用Token中的真实userId', '', ''],
    ['E-012', '重复提交订单', '快速连续两次提交相同订单', '只创建一个订单', '', ''],
    ['E-013', '并发扣库存', '多用户同时购买最后一件库存', '只有一个用户成功，其他提示库存不足', '', ''],
    ['E-014', '超大请求体', '发送超大JSON请求', '返回400或413，不导致服务崩溃', '', ''],
])

import os
output = '功能测试文档_更新版.docx'
doc.save(output)
print('Done! Saved as ' + output)
